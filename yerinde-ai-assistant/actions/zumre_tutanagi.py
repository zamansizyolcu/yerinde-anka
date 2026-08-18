"""
actions/zumre_tutanagi.py — Zümre toplantı tutanağı üretimi.

Akış:
  1) Kullanıcı önce "DOSYA YÜKLE" ile okulunun/zümresinin GERÇEK kullandığı bir
     tutanak örneğini (.docx) yükler.
  2) Bu modül örneği python-docx ile açar; okul/zümre başlığını, bilgi
     tablosunu (Toplantı No/Yeri/Tarihi/Zümre Başkanı...) ve GÜNDEM MADDESİ /
     ALINAN KARARLAR tablosunu (MEB'in resmî tutanaklarda tablo çoklu sayfaya
     yayıldığında birden fazla ayrı <w:tbl> olarak böldüğü biçim de dahil)
     ayrıştırır.
  3) Gündem maddelerinin METNİ ASLA DEĞİŞTİRİLMEZ — bunlar MEB Eğitim
     Kurulları ve Zümreleri Yönergesi Madde-12 kaynaklı, sabit resmî
     ifadelerdir. Sadece "Alınan Kararlar" sütunu, aktif model (Gemini ya da
     Ollama — hangisi kullanılıyorsa) ile YENİ dönem/tarihe uygun olarak
     yeniden üretilir.
  4) Sonuç, örnek dosyanın TÜM biçimlendirmesi (font, tablo kenarlıkları,
     kalın başlıklar, sayfa düzeni) korunarak yeni bir .docx olarak
     Masaüstü/Çalışmalarım/Zümre Tutanakları klasörüne kaydedilir ve açılır.

Hem Gemini hem Ollama (ör. qwen3:30b-a3b) ile çalışır — hangisi o an aktifse
onunla üretir; ekstra bir ayar gerekmez.
"""

from __future__ import annotations

import json
import platform
import re
import subprocess
import time
from datetime import datetime
from pathlib import Path

import requests

from app_config import get_app_config_value, get_model_provider, ollama_think_value

_IS_WINDOWS = platform.system() == "Windows"


def _tr_lower(s: str) -> str:
    """Türkçe İ/I çiftini doğru küçültür (Python'un varsayılan .lower()'ı
    ASCII 'I' harfini yanlışlıkla noktalı 'i' yapar, 'ı' değil)."""
    return s.replace("İ", "i").replace("I", "ı").lower()

# Bir LLM çağrısında en fazla bu kadar gündem maddesi için karar üretilir —
# yerel modellerde (özellikle Qwen3-30B-A3B gibi büyük MoE modellerde,
# zayıf/orta donanımda) tek seferde 30+ madde istemek hem çok uzun sürer hem
# de modelin JSON'u yarıda kesme riskini artırır. Küçük gruplar hem daha
# güvenilir hem de bir grup başarısız olursa sadece o grubu etkiler.
BATCH_SIZE = 6

# Bir grup HTTP isteğinin azami bekleme süresi. Qwen3-30B-A3B gibi büyük
# modeller zayıf donanımda yavaş olabilir; sesli komutlardaki kısa
# timeout'ların (120sn) aksine burada bilerek çok daha cömert davranıyoruz —
# bu, canlı ses döngüsünü BLOKE ETMEYEN, ayrı/tek seferlik bir araç çağrısı
# olduğu için projenin geri kalanına performans yükü bindirmez.
LLM_TIMEOUT = 600


# ══ 1) Yardımcılar ═══════════════════════════════════════════════════════
def _resolve_reference(file_path: str, ders: str = "") -> Path | None:
    candidate = (file_path or "").strip()
    if not candidate and ders:
        from actions.belge_referanslari import referans_yolu
        key = "robotik_zumre_ornek" if "robotik" in ders.lower() else "bilisim_zumre_ornek"
        p = referans_yolu(key)
        if p:
            return p
    if not candidate:
        candidate = str(get_app_config_value("last_uploaded_file", "") or "").strip()
    if not candidate:
        return None
    p = Path(candidate)
    return p if p.exists() else None


def _open_file(path: Path) -> None:
    if _IS_WINDOWS:
        import os
        os.startfile(str(path))  # noqa: S606
    else:
        subprocess.Popen(["xdg-open", str(path)],
                         stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)


def _set_paragraph_text(paragraph, new_text: str) -> None:
    """Bir paragrafın metnini, ilk run'ın biçimlendirmesini koruyarak değiştirir."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def _set_cell_text(cell, new_text: str) -> None:
    """Bir tablo hücresinin metnini, hücrenin/ilk run'ın biçimlendirmesini
    (font, kalınlık vb.) BOZMADAN değiştirir. python-docx'te `cell.text = x`
    kullanmak biçimlendirmeyi sıfırlar; bunun yerine ilk paragrafın ilk
    run'ını koruyup metnini değiştiriyor, fazla paragraf/run'ları temizliyoruz."""
    paragraphs = cell.paragraphs
    # Fazla paragrafları sil, sadece ilkini bırak
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)
    _set_paragraph_text(cell.paragraphs[0], new_text)


_DERS_YILI_RE = re.compile(r"\b(\d{4})\s*[-–]\s*(\d{4})\b")
_DONEM_MENTION_RE = re.compile(
    r"(SENE\s+SONU|DÖNEM\s+SONU|DÖNEM\s+BA[ŞS]I|DÖNEM\s+ORTASI|ARA\s+TOPLANTI)",
    re.IGNORECASE,
)
# Python'un genel .title()/.capitalize() metodları Türkçe İ/I çiftinde hataya
# düşer (ör. "BAŞI".title() -> yanlışlıkla "Başi" verir, "Başı" değil). Bu
# yüzden her iki büyük/küçük harf biçimini SABİT olarak tanımlıyoruz.
_DONEM_FORMS = {
    "dönem başı":   {"upper": "DÖNEM BAŞI",   "title": "Dönem Başı"},
    "dönem ortası": {"upper": "DÖNEM ORTASI", "title": "Dönem Ortası"},
    "ara toplantı": {"upper": "ARA TOPLANTI", "title": "Ara Toplantı"},
    "dönem sonu":   {"upper": "DÖNEM SONU",   "title": "Dönem Sonu"},
}


def _replace_donem_mentions(text: str, yeni_donem_turu: str) -> str:
    forms = _DONEM_FORMS.get(yeni_donem_turu, _DONEM_FORMS["dönem sonu"])

    def _sub(m):
        return forms["upper"] if m.group(0).isupper() else forms["title"]

    return _DONEM_MENTION_RE.sub(_sub, text)


def _apply_title_and_paren_updates(doc, ctx: dict) -> None:
    """Başlıktaki ders yılı / dönem türü ibarelerini ve gündem tablosu
    üst bilgisindeki '(Dönem Sonu)' benzeri parantezli etiketleri günceller.
    Gündem maddesi / karar METİNLERİNE ASLA dokunmaz — sadece başlık
    paragrafları ve 'Gündem Maddesi' ile başlayan üst bilgi satırları."""
    # Gövde paragrafları — SADECE ORTALANMIŞ (başlık bloğu) paragraflar
    # işlenir. Madde-12 yönerge alıntısı gibi düz metinler ORTALANMAMIŞ
    # olduğundan dokunulmaz (aksi hâlde "...ikinci dönem başında..." gibi
    # normal bir cümle içindeki kelimeler yanlışlıkla değiştirilebilir).
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for p in doc.paragraphs:
        if p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            continue
        original = p.text
        if not original.strip():
            continue
        updated = _DERS_YILI_RE.sub(ctx["ders_yili"], original)
        updated = _replace_donem_mentions(updated, ctx["donem_turu"])
        if updated != original:
            _set_paragraph_text(p, updated)

    # Tablo üst bilgi satırları — bilgi tablosu (Toplantı No vb., ör. "2026/4
    # (Sene Sonu)") ve "Gündem Maddesi" ile başlayan gündem üst bilgi satırları
    for ti, t in enumerate(doc.tables):
        for row in t.rows:
            if not row.cells:
                continue
            is_info_table = (ti == 0)
            is_header_row = row.cells[0].text.strip().startswith("Gündem Maddesi")
            if not (is_info_table or is_header_row):
                continue
            for cell in row.cells:
                for p in cell.paragraphs:
                    original = p.text
                    updated = _replace_donem_mentions(original, ctx["donem_turu"])
                    if updated != original:
                        _set_paragraph_text(p, updated)


def _next_donem_takvimi(bugun: datetime | None = None) -> dict:
    """Bugünün tarihine göre makul bir varsayılan ders yılı / dönem türü tahmini."""
    d = bugun or datetime.now()
    if d.month >= 9:  # Eylül-Aralık: yeni ders yılı başlangıcı
        ders_yili = f"{d.year}-{d.year + 1}"
        donem = "dönem başı"
    elif d.month <= 1:  # Ocak: birinci dönem sonu / ikinci dönem başı civarı
        ders_yili = f"{d.year - 1}-{d.year}"
        donem = "dönem ortası"
    elif d.month <= 6:  # Şubat-Haziran: ikinci dönem / sene sonu
        ders_yili = f"{d.year - 1}-{d.year}"
        donem = "dönem sonu"
    else:  # Temmuz-Ağustos: tatil, bir sonraki döneme hazırlık
        ders_yili = f"{d.year}-{d.year + 1}"
        donem = "dönem başı"
    return {"ders_yili": ders_yili, "donem_turu": donem}


# ══ 2) Örnek dosyayı ayrıştırma ═════════════════════════════════════════════
def _read_reference(path: Path) -> dict:
    import docx
    d = docx.Document(str(path))

    # Başlık paragraflarından zümre adını çıkar (ör. "...BİLİŞİM TEKNOLOJİLERİ
    # ZÜMRESİ..." -> "BİLİŞİM TEKNOLOJİLERİ"). İlk ~6 paragrafa bakılır.
    title_text = " ".join(p.text for p in d.paragraphs[:8] if p.text.strip())
    m = re.search(r"YILI\s+(.+?)\s+ZÜMRESİ", title_text, re.S)
    zumre_adi = m.group(1).strip() if m else ""
    donem_hint = "sene sonu" if "SENE SONU" in title_text.upper() else (
        "dönem başı" if "DÖNEM BAŞI" in title_text.upper() else "")

    if len(d.tables) < 2:
        raise ValueError(
            "Bu dosyada beklenen tabloları bulamadım (bilgi tablosu + gündem/"
            "karar tablosu). Gerçek bir zümre tutanağı örneği mi yükledin?")

    info_table = d.tables[0]
    info = {}
    for row in info_table.rows:
        if len(row.cells) < 2:
            continue
        label = row.cells[0].text.strip().rstrip(":")
        value = row.cells[1].text.strip()
        if label:
            info[label] = value

    # Gündem/karar: tables[1:] içindeki, başlık satırı OLMAYAN tüm satırlar.
    gundem_rows = []  # (tablo_index, satir_index, gundem_metni, eski_karar)
    for ti, t in enumerate(d.tables[1:], start=1):
        for ri, row in enumerate(t.rows):
            if len(row.cells) < 2:
                continue
            left = row.cells[0].text.strip()
            right = row.cells[1].text.strip()
            if left.startswith("Gündem Maddesi") or not left:
                continue
            gundem_rows.append((ti, ri, left, right))

    if not gundem_rows:
        raise ValueError("Gündem maddesi / karar satırı bulunamadı — dosya formatı beklenenden farklı olabilir.")

    return {
        "docx": d,
        "path": path,
        "zumre_adi": zumre_adi,
        "donem_hint": donem_hint,
        "info": info,
        "gundem_rows": gundem_rows,
    }


# ══ 3) LLM ile yeni kararlar üretme ══════════════════════════════════════
def _prompt_for_batch(gundem_maddeleri: list[str], ctx: dict) -> str:
    numbered = "\n".join(f"{i+1}. {g}" for i, g in enumerate(gundem_maddeleri))
    ek = f"\nEK TALİMAT: {ctx['ek_talimat']}" if ctx.get("ek_talimat") else ""
    return (
        f"Sen bir Türk devlet ortaokulunda '{ctx['zumre_adi']}' zümresi için resmî "
        f"zümre toplantı tutanağı yazan bir zümre başkanısın. Aşağıda, MEB Eğitim "
        f"Kurulları ve Zümreleri Yönergesi'nden gelen SABİT gündem maddeleri var. "
        f"{ctx['ders_yili']} eğitim-öğretim yılı, {ctx['donem_turu']} toplantısı "
        f"({ctx['toplanti_tarihi']} tarihli) için, HER gündem maddesine karşılık "
        f"resmî/profesyonel Türkçe, üçüncü tekil/çoğul şahıs, tutanak üslubunda "
        f"2-4 cümlelik bir 'Alınan Kararlar ve Değerlendirmeler' metni yaz.\n\n"
        f"ÖNEMLİ ÜSLUP KURALI: Toplantı türü '{ctx['donem_turu']}' ise —\n"
        f"  - 'dönem başı' ya da 'dönem ortası' ise: GELECEĞE yönelik PLANLAMA "
        f"dilinde yaz (\"...planlanmıştır\", \"...yapılacaktır\", \"...hedeflenmektedir\").\n"
        f"  - 'dönem sonu' ise: GEÇMİŞE yönelik DEĞERLENDİRME dilinde yaz "
        f"(\"...başarıyla tamamlanmıştır\", \"...gerçekleştirilmiştir\").\n"
        f"{ek}\n\n"
        f"GÜNDEM MADDELERİ:\n{numbered}\n\n"
        f"SADECE şu JSON dizisini döndür, başka HİÇBİR şey yazma (açıklama, "
        f"markdown kod bloğu (```), giriş cümlesi YOK): "
        f'["1. maddenin kararı", "2. maddenin kararı", ...] '
        f"— tam olarak {len(gundem_maddeleri)} elemanlı olmalı, sıra korunmalı."
    )


def _extract_json_array(raw: str, expected_len: int) -> list[str] | None:
    raw = raw.strip()
    # Markdown kod bloğu içine sarılmışsa temizle
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw.strip())
    m = re.search(r"\[.*\]", raw, re.S)
    if not m:
        return None
    try:
        arr = json.loads(m.group(0))
    except Exception:
        return None
    if not isinstance(arr, list) or len(arr) != expected_len:
        return None
    return [str(x).strip() for x in arr]


def _call_gemini_text(prompt: str) -> str | None:
    api_key = str(get_app_config_value("gemini_api_key", "") or "").strip()
    if not api_key:
        return None
    try:
        from google import genai
    except ImportError:
        return None
    try:
        client = genai.Client(api_key=api_key)
        for model in ("models/gemini-2.5-flash", "models/gemini-2.0-flash", "models/gemini-2.5-flash-lite"):
            try:
                resp = client.models.generate_content(model=model, contents=[prompt])
                if resp and resp.text:
                    return resp.text
            except Exception:
                continue
    except Exception:
        pass
    return None


def _call_ollama_text(prompt: str) -> str | None:
    try:
        host = str(get_app_config_value("ollama_host", "http://localhost:11434") or "http://localhost:11434")
        model = str(get_app_config_value("ollama_model", "llama3.1") or "llama3.1")
        r = requests.post(f"{host}/api/chat", json={
            "model": model, "stream": False,
            "messages": [{"role": "user", "content": prompt}],
            "options": {"temperature": 0.4, "num_predict": 2200},
            "think": ollama_think_value(),
            "keep_alive": "30m",
        }, timeout=LLM_TIMEOUT)
        if r.status_code >= 400:
            # 'think' alanı bu modelde desteklenmiyor olabilir — kaldırıp tekrar dene.
            payload2 = {"model": model, "stream": False,
                        "messages": [{"role": "user", "content": prompt}],
                        "options": {"temperature": 0.4, "num_predict": 2200},
                        "keep_alive": "30m"}
            r = requests.post(f"{host}/api/chat", json=payload2, timeout=LLM_TIMEOUT)
        return r.json().get("message", {}).get("content", "")
    except Exception:
        return None


def _generate_batch(gundem_maddeleri: list[str], ctx: dict) -> list[str]:
    prompt = _prompt_for_batch(gundem_maddeleri, ctx)
    provider = get_model_provider()
    raw = _call_gemini_text(prompt) if provider == "gemini" else _call_ollama_text(prompt)
    if raw:
        arr = _extract_json_array(raw, len(gundem_maddeleri))
        if arr:
            return arr
        # tek seferlik, daha katı bir tekrar deneme
        raw2 = _call_gemini_text(prompt + "\n\nUNUTMA: SADECE JSON dizisi.") if provider == "gemini" \
            else _call_ollama_text(prompt + "\n\nUNUTMA: SADECE JSON dizisi.")
        if raw2:
            arr2 = _extract_json_array(raw2, len(gundem_maddeleri))
            if arr2:
                return arr2
    # Son çare: boş bırakmak yerine genel ama dönem türüne uygun tek cümlelik
    # bir yer tutucu üret — kullanıcı sonradan elle düzenleyebilir.
    genel = ("Bu madde kapsamında gerekli çalışmalar planlanmıştır."
             if ctx["donem_turu"] in ("dönem başı", "dönem ortası")
             else "Bu madde kapsamında gerekli çalışmalar tamamlanmıştır.")
    return [genel] * len(gundem_maddeleri)


def _generate_all_kararlar(gundem_rows: list[tuple], ctx: dict) -> tuple[list[str], int]:
    """Tüm gündem maddeleri için karar üretir (gruplar hâlinde). İkinci
    dönüş değeri, yer tutucuya düşen madde sayısıdır (kullanıcıyı uyarmak için)."""
    maddeler = [g[2] for g in gundem_rows]
    sonuc: list[str] = []
    fallback_count = 0
    genel_fallbacks = {
        "Bu madde kapsamında gerekli çalışmalar planlanmıştır.",
        "Bu madde kapsamında gerekli çalışmalar tamamlanmıştır.",
    }
    for i in range(0, len(maddeler), BATCH_SIZE):
        batch = maddeler[i:i + BATCH_SIZE]
        kararlar = _generate_batch(batch, ctx)
        fallback_count += sum(1 for k in kararlar if k in genel_fallbacks)
        sonuc.extend(kararlar)
    return sonuc, fallback_count


# ══ 4) Yeni dosyayı üretme ═══════════════════════════════════════════════
def zumre_tutanagi_olustur(
    donem_turu: str = "",
    toplanti_tarihi: str = "",
    toplanti_saati: str = "",
    toplanti_no: str = "",
    ders_yili: str = "",
    ek_talimat: str = "",
    dosya_yolu: str = "",
    ders: str = "",
) -> str:
    ref_path = _resolve_reference(dosya_yolu, ders)
    if not ref_path:
        return ("Örnek bir zümre tutanağı (.docx) bulamadım. Önce 'DOSYA YÜKLE' "
                "düğmesiyle okulunun kullandığı gerçek bir tutanak örneğini yükle "
                "(istersen 'bunu bilişim/robotik zümre örneği olarak kaydet' de ki "
                "bir daha yüklemene gerek kalmasın), sonra tekrar iste.")

    try:
        import docx  # noqa: F401
    except ImportError:
        return "Word dosyalarıyla çalışmak için 'pip install python-docx' gerekli."

    try:
        ref = _read_reference(ref_path)
    except Exception as e:
        return f"Örnek dosya okunamadı: {e}"

    varsayilan = _next_donem_takvimi()
    ctx = {
        "zumre_adi": ref["zumre_adi"] or "İlgili",
        "donem_turu": _tr_lower((donem_turu or varsayilan["donem_turu"]).strip()),
        "ders_yili": (ders_yili or varsayilan["ders_yili"]).strip(),
        "toplanti_tarihi": (toplanti_tarihi or datetime.now().strftime("%d.%m.%Y")).strip(),
        "ek_talimat": ek_talimat.strip(),
    }

    n_madde = len(ref["gundem_rows"])
    provider = get_model_provider()
    model_adi = (str(get_app_config_value("ollama_model", "")) if provider == "ollama"
                 else "Gemini")

    yeni_kararlar, fallback_count = _generate_all_kararlar(ref["gundem_rows"], ctx)

    # ── Bilgi tablosunu güncelle ─────────────────────────────────────────
    info_table = ref["docx"].tables[0]
    for row in info_table.rows:
        if len(row.cells) < 2:
            continue
        label = row.cells[0].text.strip().rstrip(":")
        if label.startswith("Toplantı No") and toplanti_no:
            _set_cell_text(row.cells[1], toplanti_no)
        elif label.startswith("Toplantı Tarihi"):
            saat = toplanti_saati.strip() or row.cells[1].text.split("–")[-1].strip()
            _set_cell_text(row.cells[1], f"{ctx['toplanti_tarihi']} – {saat}")

    # ── Gündem/karar hücrelerini güncelle ────────────────────────────────
    tables = ref["docx"].tables
    for (ti, ri, _gundem, _eski), yeni in zip(ref["gundem_rows"], yeni_kararlar):
        cell = tables[ti].rows[ri].cells[1]
        _set_cell_text(cell, yeni)

    # ── OLUR tarihini güncelle (varsa) ───────────────────────────────────
    _DATE_RE = re.compile(r"^\d{2}[./]\d{2}[./]\d{4}$")
    for p in ref["docx"].paragraphs:
        if _DATE_RE.match(p.text.strip()):
            _set_paragraph_text(p, ctx["toplanti_tarihi"])

    # ── Başlıktaki ders yılı / dönem türü ibarelerini güncelle ────────────
    _apply_title_and_paren_updates(ref["docx"], ctx)

    # ── Kaydet ────────────────────────────────────────────────────────────
    from actions.code_tools import ensure_workspace_folder
    folder = ensure_workspace_folder("Zümre Tutanakları")
    stamp = time.strftime("%Y-%m-%d %H.%M")
    donem_etiketi = ctx["donem_turu"].replace(" ", "-")
    out_name = f"Zümre Tutanağı - {ref['zumre_adi']} - {donem_etiketi} {stamp}.docx"
    out_path = folder / out_name
    i = 2
    while out_path.exists():
        out_path = folder / f"Zümre Tutanağı - {ref['zumre_adi']} - {donem_etiketi} {stamp} ({i}).docx"
        i += 1

    try:
        ref["docx"].save(str(out_path))
    except Exception as e:
        return f"Dosya kaydedilemedi: {e}"

    try:
        _open_file(out_path)
    except Exception:
        pass

    uyari = ""
    if fallback_count:
        uyari = (f" NOT: {fallback_count} madde için model yanıtı ayrıştırılamadığından "
                 f"genel bir yer tutucu metin kullanıldı — bu maddeleri açıp elle "
                 f"gözden geçirmeni öneririm.")

    return (f"'{ref['zumre_adi']}' zümresi için {ctx['donem_turu']} tutanağı hazırlandı "
            f"({n_madde} gündem maddesi, {model_adi} ile üretildi) ve açıldı: "
            f"{out_path.name}.{uyari}")
