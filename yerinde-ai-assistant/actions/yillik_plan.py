"""
actions/yillik_plan.py — Yıllık plandan haftalık "günlük plan" üretimi.

GERÇEK VERİYLE DOĞRULANMIŞ ÖNEMLİ NOT: Farklı sınıfların/derslerin yıllık plan
Excel dosyaları BİRBİRİNDEN FARKLI sütun düzenlerinde geliyor (ör. 5. sınıf
Bilişim planı "TDE" / yeni müfredat düzeninde BTY.5.x.x kodlarıyla 8 sütun,
6. sınıf aynı dersin planı ise klasik AY|HAFTA|GÜN|SAAT|KAZANIMLAR 5 sütun
düzeninde BT.6.x.x.x kodlarıyla). Bu yüzden sabit sütun indeksi varsaymak
yerine, başlık satırındaki metinlerden ESNEK bir sütun eşlemesi çıkarıyoruz.

Akış:
  1) Yıllık plan (.xlsx) satır satır taranır; "ÜNİTE:" / "ÖĞRENME ALANI"
     içeren satırlar ünite başlığı olarak yakalanıp bir sonraki asıl veri
     satırlarına forward-fill edilir (aynı şekilde AY sütunu da).
  2) İstenen hafta (numarayla ya da konu metniyle aranarak) bulunur.
  3) O haftanın kazanım/konu bilgisi bağlam olarak modele verilir; model
     sadece DERS İŞLENİŞ adımlarını, materyalleri ve ölçme-değerlendirmeyi
     üretir — kazanım kodlarının METNİNİ ASLA DEĞİŞTİRMEZ.
  4) Örnek günlük plan biçimine (açık mavi etiketli 2 sütunlu tablo) uygun
     yeni bir .docx python-docx ile inşa edilir.
"""

from __future__ import annotations

import json
import platform
import re
import subprocess
import time
from pathlib import Path

import requests

from app_config import get_app_config_value, get_model_provider, ollama_think_value

_IS_WINDOWS = platform.system() == "Windows"
LLM_TIMEOUT = 600

DERS_BILGI = {
    "bilişim": {"ad": "BİLİŞİM TEKNOLOJİLERİ VE YAZILIM DERSİ", "brans": "Bilişim Teknolojileri ve Yazılım Öğretmeni"},
    "robotik": {"ad": "SEÇMELİ ROBOTİK KODLAMA DERSİ", "brans": "Bilişim Teknolojileri ve Yazılım Öğretmeni"},
}


def _tr_lower(s: str) -> str:
    return s.replace("İ", "i").replace("I", "ı").lower()


def _normalize_ders(ders: str) -> str:
    d = _tr_lower(ders or "")
    return "robotik" if "robot" in d else "bilişim"


def _open_file(path: Path) -> None:
    if _IS_WINDOWS:
        import os
        os.startfile(str(path))  # noqa: S606
    else:
        subprocess.Popen(["xdg-open", str(path)],
                         stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)


# ══ 1) Yıllık planı esnek biçimde ayrıştırma ═══════════════════════════════
_KONSEPT_ANAHTAR = {
    "ay": ["ay"],
    "hafta": ["hafta"],
    "gun": ["gün"],
    "sure": ["saat", "süre"],
    "unite": ["ünite", "tema"],
    "kazanim": ["kazanım", "öğrenme çıktı", "öğrenme çiktilari"],
    "konu": ["konu"],
    "olcme": ["ölçme", "öğrenme kanit", "değerlendirme"],
    "materyal": ["materyal", "araç"],
}


def _bul_baslik_satiri(rows: list[tuple]) -> int | None:
    for i, row in enumerate(rows[:12]):
        cells = [_tr_lower(str(c or "")) for c in row]
        if any("hafta" == c.strip() or "hafta" in c for c in cells) and \
           any("ay" == c.strip() for c in cells):
            return i
    return None


def _sutun_esle(header_row: tuple) -> dict:
    colmap = {}
    for idx, cell in enumerate(header_row):
        text = _tr_lower(str(cell or ""))
        if not text.strip():
            continue
        for kavram, anahtarlar in _KONSEPT_ANAHTAR.items():
            if kavram in colmap:
                continue
            if any(a in text for a in anahtarlar):
                colmap[kavram] = idx
    return colmap


def _load_yillik_plan(path: Path) -> list[dict]:
    """Döner: [{ay, hafta_no, hafta_ham, sure, unite, kazanim, konu, olcme}, ...]"""
    import openpyxl
    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    haftalar: list[dict] = []
    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        hdr_idx = _bul_baslik_satiri(rows)
        if hdr_idx is None:
            continue
        colmap = _sutun_esle(rows[hdr_idx])
        if "hafta" not in colmap:
            continue

        mevcut_ay = ""
        mevcut_unite = ""

        for row in rows[hdr_idx + 1:]:
            cells = list(row)
            dolu = [c for c in cells if c not in (None, "")]
            if not dolu:
                continue
            satir_metni = " ".join(str(c) for c in dolu)

            # AY bilgisini, satır ünite-duyurusu OLSA BİLE önce yakala (bazı
            # düzenlerde "EYLÜL" ile ünite duyurusu aynı satırda geliyor).
            ay_val = cells[colmap["ay"]] if "ay" in colmap and colmap["ay"] < len(cells) else None
            if ay_val and str(ay_val).strip():
                mevcut_ay = str(ay_val).strip()

            # Ünite/öğrenme alanı duyuru satırı: az sayıda dolu hücre + "ÜNİTE"/"ÖĞRENME ALANI" geçiyor
            if len(dolu) <= 2 and re.search(r"ÜNİTE|ÖĞRENME ALANI", satir_metni, re.IGNORECASE):
                unite_hucreleri = [str(c) for idx, c in enumerate(cells)
                                    if c not in (None, "") and idx != colmap.get("ay", -1)]
                mevcut_unite = " ".join(unite_hucreleri).strip() or satir_metni.strip()
                continue

            hafta_val = cells[colmap["hafta"]] if colmap["hafta"] < len(cells) else None
            if hafta_val is None or not str(hafta_val).strip():
                continue
            hafta_ham = str(hafta_val).strip()
            m = re.search(r"\d+", hafta_ham)
            if not m:
                # "1. DÖNEM ARA TATİLİ" gibi tatil/bilgi satırları — hafta değil, atla
                continue
            hafta_no = int(m.group(0))

            def _al(kavram):
                i = colmap.get(kavram)
                if i is None or i >= len(cells) or cells[i] is None:
                    return ""
                return str(cells[i]).strip()

            unite_bu_satir = _al("unite")
            if unite_bu_satir:
                mevcut_unite = unite_bu_satir

            kazanim = _al("kazanim")
            if not kazanim:
                # bazı düzenlerde kazanım ayrı sütunda değil, ünite satırının
                # kendisinde olabilir — en azından boş bırakmayalım
                kazanim = unite_bu_satir

            haftalar.append({
                "ay": mevcut_ay, "hafta_no": hafta_no, "hafta_ham": hafta_ham,
                "sure": _al("sure"), "unite": mevcut_unite,
                "kazanim": kazanim, "konu": _al("konu"), "olcme": _al("olcme"),
            })
    return haftalar


def _select_hafta(haftalar: list[dict], hafta_no: int | None, konu_arama: str) -> dict | None:
    if hafta_no is not None:
        for h in haftalar:
            if h["hafta_no"] == hafta_no:
                return h
        return None
    if konu_arama:
        arama = _tr_lower(konu_arama)
        for h in haftalar:
            metin = _tr_lower(h["kazanim"] + " " + h["unite"] + " " + h["konu"])
            if arama in metin:
                return h
    return None


# ══ 2) LLM ile ders işleniş / materyal / ölçme üretimi ═══════════════════
def _extract_json_obj(raw: str) -> dict | None:
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw.strip())
    m = re.search(r"\{.*\}", raw, re.S)
    if not m:
        return None
    try:
        return json.loads(m.group(0))
    except Exception:
        return None


def _call_gemini_text(prompt: str) -> str | None:
    api_key = str(get_app_config_value("gemini_api_key", "") or "").strip()
    if not api_key:
        return None
    try:
        from google import genai
    except ImportError:
        return None
    try:
        client = genai.Client(api_key=api_key)
        for model in ("models/gemini-2.5-flash", "models/gemini-2.0-flash", "models/gemini-2.5-flash-lite"):
            try:
                resp = client.models.generate_content(model=model, contents=[prompt])
                if resp and resp.text:
                    return resp.text
            except Exception:
                continue
    except Exception:
        pass
    return None


def _call_ollama_text(prompt: str, num_predict: int = 1600) -> str | None:
    try:
        host = str(get_app_config_value("ollama_host", "http://localhost:11434") or "http://localhost:11434")
        model = str(get_app_config_value("ollama_model", "llama3.1") or "llama3.1")
        payload = {"model": model, "stream": False,
                   "messages": [{"role": "user", "content": prompt}],
                   "options": {"temperature": 0.5, "num_predict": num_predict},
                   "think": ollama_think_value(), "keep_alive": "30m"}
        r = requests.post(f"{host}/api/chat", json=payload, timeout=LLM_TIMEOUT)
        if r.status_code >= 400:
            payload.pop("think", None)
            r = requests.post(f"{host}/api/chat", json=payload, timeout=LLM_TIMEOUT)
        return r.json().get("message", {}).get("content", "")
    except Exception:
        return None


def _call_llm(prompt: str, num_predict: int = 1600) -> str | None:
    return (_call_gemini_text(prompt) if get_model_provider() == "gemini"
            else _call_ollama_text(prompt, num_predict))


def _icerik_uret(hafta: dict, ctx: dict) -> dict:
    ek = f"\nEK TALİMAT: {ctx['ek_talimat']}" if ctx.get("ek_talimat") else ""
    prompt = (
        f"Sen bir Türk devlet ortaokulunda {ctx['sinif']}. sınıflara {ctx['ders_ad']} "
        f"dersi veren bir öğretmensin. Aşağıda YILLIK PLANINDAN bir haftanın "
        f"kazanım/konu bilgisi var. Bu bilgiye dayanarak o haftanın GÜNLÜK "
        f"(HAFTALIK) DERS PLANINI hazırla.{ek}\n\n"
        f"ÜNİTE: {hafta['unite'] or '(belirtilmemiş)'}\n"
        f"KAZANIM/KONU (yıllık plandan, AYNEN KORU, DEĞİŞTİRME): {hafta['kazanim']}\n"
        f"{'EK KONU BAŞLIĞI: ' + hafta['konu'] if hafta['konu'] else ''}\n"
        f"{'YILLIK PLANDAKİ ÖLÇME/KANIT ÖNERİSİ: ' + hafta['olcme'] if hafta['olcme'] else ''}\n\n"
        f"SADECE şu JSON nesnesini döndür, başka HİÇBİR şey yazma:\n"
        '{"konu_basligi": "kısa, öğretmen tarzı 3-6 kelimelik konu başlığı", '
        '"materyaller": "kullanılacak araç-gereç/dijital materyaller, kısa liste, virgülle ayrılmış", '
        '"ders_isleyisi": ["adım 1", "adım 2", "..."], '
        '"olcme_degerlendirme": "1-2 cümlelik ölçme-değerlendirme yöntemi"}\n'
        f"ders_isleyisi 6-10 arası, numaralanmamış (numarayı biz ekleyeceğiz), "
        f"somut ve {ctx['sinif']}. sınıf seviyesine uygun sınıf-içi eylem cümleleri olsun "
        f"(ör. '... videosu izletilir', '... etkinliği yapılır', '... sorusu sorulur')."
    )
    raw = _call_llm(prompt)
    obj = _extract_json_obj(raw) if raw else None
    if not obj:
        raw2 = _call_llm(prompt + "\n\nUNUTMA: SADECE JSON nesnesi.") if raw else None
        obj = _extract_json_obj(raw2) if raw2 else None
    if not obj:
        return {
            "konu_basligi": hafta["konu"] or hafta["unite"] or "Konu",
            "materyaller": "Akıllı Tahta, ilgili sunum/video",
            "ders_isleyisi": ["(Model yanıt vermedi — bu bölümü elle doldurun.)"],
            "olcme_degerlendirme": "Konuyla ilgili bir çalışma kağıdı/etkinlik ile değerlendirme yapılır.",
        }
    obj.setdefault("konu_basligi", hafta["konu"] or hafta["unite"] or "Konu")
    obj.setdefault("materyaller", "Akıllı Tahta")
    obj.setdefault("ders_isleyisi", [])
    obj.setdefault("olcme_degerlendirme", "")
    return obj


# ══ 3) Word belgesini inşa etme ═══════════════════════════════════════════
def _build_plan_docx(ctx: dict, hafta: dict, icerik: dict, out_path: Path) -> None:
    import docx
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    d = docx.Document()
    section = d.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Cm(1.5))

    def _shade(cell, hex_renk="BDD7EE"):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), hex_renk)
        tcPr.append(shd)

    okul = ctx.get("okul_adi") or "[OKUL ADI]"
    p1 = d.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run(okul)
    r.bold = True
    r.font.size = Pt(13)

    p2 = d.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run(f"{ctx['ders_ad']} GÜNLÜK PLAN")
    r.bold = True
    r.font.size = Pt(13)
    p2.paragraph_format.space_after = Pt(12)

    satirlar = [
        ("SINIF:", f"{ctx['sinif']}. SINIFLAR"),
        ("ÜNİTE:", hafta["unite"] or "-"),
        ("KONU:", icerik["konu_basligi"]),
        ("HEDEFLER:", hafta["kazanim"] or "-"),
        ("TARİH:", f"{hafta['hafta_no']}.HAFTA" + (f" ({hafta['ay']})" if hafta.get("ay") else "")),
        ("SÜRE:", (hafta["sure"] or "2") if "SAAT" in str(hafta.get("sure", "")).upper()
                  else f"{hafta['sure'] or '2'} DERS SAATİ"),
        ("KULLANILAN\nMATERYALLER:", icerik["materyaller"]),
        ("DERS İŞLENİŞ :", None),  # özel işlenecek (numaralı liste)
        ("ÖLÇME\nDEĞERLENDİRME :", icerik["olcme_degerlendirme"]),
    ]

    tbl = d.add_table(rows=len(satirlar), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(4.2)
    tbl.columns[1].width = Cm(13.5)

    for i, (etiket, deger) in enumerate(satirlar):
        lcell = tbl.rows[i].cells[0]
        lcell.text = ""
        lp = lcell.paragraphs[0]
        for satir in etiket.split("\n"):
            if lp.runs or lp.text:
                lp = lcell.add_paragraph()
            rr = lp.add_run(satir)
            rr.bold = True
            rr.font.size = Pt(11)
        _shade(lcell, "BDD7EE")

        rcell = tbl.rows[i].cells[1]
        rcell.text = ""
        if etiket.startswith("DERS İŞLENİŞ"):
            adimlar = icerik.get("ders_isleyisi") or []
            for j, adim in enumerate(adimlar):
                p = rcell.paragraphs[0] if j == 0 else rcell.add_paragraph()
                rr = p.add_run(f"{j+1}. {adim}")
                rr.font.size = Pt(11)
        else:
            rp = rcell.paragraphs[0]
            rr = rp.add_run(str(deger or "-"))
            rr.font.size = Pt(11)

    d.save(str(out_path))


# ══ 5) Yıllık planı yeni akademik yıla uyarlama ═══════════════════════════
def _top_left(ws, row: int, col: int):
    for rng in ws.merged_cells.ranges:
        if (row, col) in rng.cells:
            return ws.cell(row=rng.min_row, column=rng.min_col)
    return ws.cell(row=row, column=col)


_TARIH_ARALIGI_RE = re.compile(
    r"(\d{1,2})\s*[-–]\s*(\d{1,2})\s*([A-ZÇĞİÖŞÜa-zçğıöşü]+)"
)
_HAFTA_ONEK_RE = re.compile(r"^(.*?Hafta\s*:?\s*\n?\s*)", re.IGNORECASE | re.DOTALL)


def yillik_plan_guncelle(
    ders: str = "bilişim",
    sinif: str = "5",
    egitim_yili: str = "",
    dosya_yolu: str = "",
) -> str:
    from actions.belge_referanslari import referans_yolu
    from actions.akademik_takvim import hafta_takvimi, bilinen_yillar

    ders_norm = _normalize_ders(ders)
    sinif = re.sub(r"\D", "", str(sinif)) or "5"
    ders_info = DERS_BILGI[ders_norm]

    if not egitim_yili:
        return ("Hangi eğitim-öğretim yılı için güncelleyeceğimi belirtir misin? "
                f"(Bildiğim akademik takvimler: {', '.join(bilinen_yillar()) or 'henüz yok'})")

    takvim = hafta_takvimi(egitim_yili)
    if not takvim:
        return (f"{egitim_yili} için akademik takvimi henüz bilmiyorum. MEB'in yayınladığı "
                f"'Çalışma Takvimi' görselini paylaşırsan (dönem başlangıç/bitiş ve ara tatil "
                f"tarihleri) ekleyip devam ederim.")

    yp_key = f"{'robotik' if ders_norm == 'robotik' else 'bilisim'}_yillik_plan"
    ref_path = (Path(dosya_yolu) if dosya_yolu and Path(dosya_yolu).exists()
                else referans_yolu(yp_key))
    if not ref_path:
        return (f"{ders_info['ad'].title()} yıllık planını bulamadım. Önce 'DOSYA YÜKLE' ile "
                f"geçen senenin planını yükleyip 'bunu {ders_norm} yıllık planı olarak kaydet' "
                f"de, sonra tekrar iste.")

    try:
        import openpyxl
    except ImportError:
        return "Excel dosyalarıyla çalışmak için 'pip install openpyxl' gerekli."

    try:
        wb = openpyxl.load_workbook(str(ref_path))
    except Exception as e:
        return f"Yıllık plan okunamadı: {e}"

    toplam_guncellenen = 0
    eksik_hafta_uyarisi = ""
    for ws in wb.worksheets:
        rows = list(ws.iter_rows())
        deger_satirlari = [[c.value for c in r] for r in rows]
        hdr_idx = _bul_baslik_satiri(deger_satirlari)
        if hdr_idx is None:
            continue
        colmap = _sutun_esle(deger_satirlari[hdr_idx])
        if "hafta" not in colmap:
            continue

        # Bu sayfadaki gerçek hafta veri satırlarını (ünite duyurusu olmayan,
        # hafta hücresinde rakam bulunan) sırayla topla.
        hafta_satirlari = []  # (excel_row_no)
        for r_idx in range(hdr_idx + 1, len(rows)):
            vals = deger_satirlari[r_idx]
            dolu = [v for v in vals if v not in (None, "")]
            if not dolu:
                continue
            satir_metni = " ".join(str(v) for v in dolu)
            if len(dolu) <= 2 and re.search(r"ÜNİTE|ÖĞRENME ALANI", satir_metni, re.IGNORECASE):
                continue
            hafta_cell_val = vals[colmap["hafta"]] if colmap["hafta"] < len(vals) else None
            if hafta_cell_val is None or not re.search(r"\d", str(hafta_cell_val)):
                continue
            hafta_satirlari.append(rows[r_idx][0].row)  # gerçek excel satır no

        yazilan_ay_hucreleri = set()

        if len(hafta_satirlari) > len(takvim):
            eksik_hafta_uyarisi = (
                f" NOT: Örnek plan {len(hafta_satirlari)} hafta içeriyor ama "
                f"{egitim_yili} takviminde {len(takvim)} hafta var — son "
                f"{len(hafta_satirlari) - len(takvim)} haftanın tarihi "
                f"güncellenemedi, elle kontrol et.")

        for pos, excel_row in enumerate(hafta_satirlari):
            if pos >= len(takvim):
                break
            hafta_bilgi = takvim[pos]
            s, e = hafta_bilgi["baslangic"], hafta_bilgi["bitis"]

            # AY sütunu (varsa) — bazı planlarda AY birden fazla hafta
            # satırını kapsayan BİRLEŞİK bir hücre (ör. A7:A11). Böyle bir
            # birleşik hücreye art arda her satır için ayrı ayrı yazarsak,
            # son yazılan satırın değeri öncekilerin üzerine geçer (yanlış
            # ay görünür). Bu yüzden her birleşik/hedef hücreye SADECE İLK
            # karşılaşıldığında yazıyoruz.
            if "ay" in colmap:
                ay_cell = ws.cell(row=excel_row, column=colmap["ay"] + 1)
                tl = _top_left(ws, ay_cell.row, ay_cell.column)
                hedef = (tl.row, tl.column)
                if hedef not in yazilan_ay_hucreleri:
                    tl.value = hafta_bilgi["ay"]
                    yazilan_ay_hucreleri.add(hedef)

            # GÜN sütunu ayrı ise (6. sınıf klasik format: "08-12"): yeni
            # gün aralığıyla değiştir.
            if "gun" in colmap:
                gun_cell = ws.cell(row=excel_row, column=colmap["gun"] + 1)
                if gun_cell.value not in (None, ""):
                    tl = _top_left(ws, gun_cell.row, gun_cell.column)
                    tl.value = f"{s.day:02d}-{e.day:02d}"

            # HAFTA hücresi — ya sade sayı (dokunma, sıra zaten korunuyor)
            # ya da "N. Hafta:\n DD-DD AyAdı" birleşik biçim (5. sınıf TDE
            # formatı) — bu durumda SADECE tarih kısmını güncelle, "N. Hafta"
            # etiketini koru.
            hafta_cell = ws.cell(row=excel_row, column=colmap["hafta"] + 1)
            eski_deger = str(hafta_cell.value or "")
            from actions.akademik_takvim import _ay_adi, ay_adi_baslik
            if s.month == e.month:
                yeni_tarih = f"{s.day}-{e.day} {ay_adi_baslik(hafta_bilgi['ay'])}"
            else:
                yeni_tarih = f"{s.day} {ay_adi_baslik(_ay_adi(s))}-{e.day} {ay_adi_baslik(_ay_adi(e))}"
            # Tarih metninin kalıbı çeşitli olabilir ("DD-DD Ay" ya da ay
            # sınırını aşan haftalarda "DD Ay-DD Ay" gibi FARKLI bir kalıp) —
            # tek bir regex'le her ihtimali eşleştirmeye çalışmak yerine,
            # "N. Hafta:" önekini koruyup ONDAN SONRAKİ HER ŞEYİ yeni tarihle
            # değiştiriyoruz; bu, kalıp ne olursa olsun güvenilir çalışır.
            m = _HAFTA_ONEK_RE.match(eski_deger)
            if m:
                tl = _top_left(ws, hafta_cell.row, hafta_cell.column)
                tl.value = m.group(1) + yeni_tarih
            elif eski_deger.strip().rstrip(".").isdigit():
                # Sade hafta numarası (ör. "1", "12.") — bu formatta tarih
                # ayrı bir GÜN sütununda tutuluyor, numaraya dokunma.
                pass
            elif re.search(r"\d", eski_deger):
                tl = _top_left(ws, hafta_cell.row, hafta_cell.column)
                tl.value = yeni_tarih

            toplam_guncellenen += 1

        # Başlık hücresindeki eski yıl ibaresini güncelle
        for row in ws.iter_rows(min_row=1, max_row=3):
            for cell in row:
                if cell.value and re.search(r"\d{4}\s*[-–]\s*\d{4}", str(cell.value)):
                    tl = _top_left(ws, cell.row, cell.column)
                    tl.value = re.sub(r"\d{4}\s*[-–]\s*\d{4}", egitim_yili, str(cell.value))

    if toplam_guncellenen == 0:
        return ("Yıllık planda güncellenecek hafta satırı bulamadım — dosya formatı "
                "beklenenden farklı olabilir.")

    from actions.code_tools import ensure_workspace_folder
    folder = ensure_workspace_folder("Yıllık Planlar")
    stamp = time.strftime("%Y-%m-%d %H.%M")
    out_name = f"{ders_norm.capitalize()} {sinif}.Sınıf Yıllık Plan {egitim_yili} {stamp}.xlsx"
    out_path = folder / out_name
    i = 2
    while out_path.exists():
        out_path = folder / f"{ders_norm.capitalize()} {sinif}.Sınıf Yıllık Plan {egitim_yili} {stamp} ({i}).xlsx"
        i += 1

    try:
        wb.save(str(out_path))
    except Exception as e:
        return f"Dosya kaydedilemedi: {e}"

    try:
        _open_file(out_path)
    except Exception:
        pass

    return (f"{ders_info['ad']} — {sinif}. sınıf yıllık planı {egitim_yili} akademik "
            f"takvimine göre güncellendi ({toplam_guncellenen} hafta tarihi yeniden "
            f"hesaplandı) ve açıldı: {out_path.name}.{eksik_hafta_uyarisi}")


# ══ 6) Yıllık plandan haftalık günlük plan üretimi ═══════════════════════
def gunluk_plan_olustur(
    ders: str = "bilişim",
    sinif: str = "5",
    hafta_no: str = "",
    konu_arama: str = "",
    ek_talimat: str = "",
    yillik_plan_dosya_yolu: str = "",
) -> str:
    from actions.belge_referanslari import referans_yolu

    ders_norm = _normalize_ders(ders)
    sinif = re.sub(r"\D", "", str(sinif)) or "5"
    ders_info = DERS_BILGI[ders_norm]

    yp_key = f"{'robotik' if ders_norm == 'robotik' else 'bilisim'}_yillik_plan"
    yp_path = (Path(yillik_plan_dosya_yolu) if yillik_plan_dosya_yolu and Path(yillik_plan_dosya_yolu).exists()
               else referans_yolu(yp_key))
    if not yp_path:
        return (f"{ders_info['ad'].title()} yıllık planını bulamadım. Önce 'DOSYA YÜKLE' ile "
                f"yükleyip 'bunu {ders_norm} yıllık planı olarak kaydet' de, sonra tekrar iste.")

    try:
        haftalar = _load_yillik_plan(yp_path)
    except Exception as e:
        return f"Yıllık plan okunamadı: {e}"
    if not haftalar:
        return ("Yıllık plandan hafta bilgisi çıkaramadım — dosya formatı beklenenden farklı "
                "olabilir. Dosyayı bana biraz daha anlatır mısın?")

    hafta_no_int = int(re.sub(r"\D", "", hafta_no)) if hafta_no and re.search(r"\d", hafta_no) else None
    hafta = _select_hafta(haftalar, hafta_no_int, konu_arama)
    if not hafta:
        if hafta_no_int is not None:
            return (f"Yıllık planda {hafta_no_int}. hafta bulunamadı (plan {len(haftalar)} hafta "
                     f"içeriyor).")
        return (f"'{konu_arama}' ile eşleşen bir hafta bulamadım. Hafta numarası verir misin "
                 f"(ör. '3. hafta için günlük plan hazırla')?")

    ctx = {"sinif": sinif, "ders_ad": ders_info["ad"], "ek_talimat": ek_talimat.strip(), "okul_adi": ""}

    # Okul adını, varsa kayıtlı zümre referansından ödünç al
    try:
        zumre_key = "robotik_zumre_ornek" if ders_norm == "robotik" else "bilisim_zumre_ornek"
        zref = referans_yolu(zumre_key)
        if zref:
            import docx as _docx
            dd = _docx.Document(str(zref))
            if dd.paragraphs and dd.paragraphs[0].text.strip():
                ctx["okul_adi"] = dd.paragraphs[0].text.strip()
    except Exception:
        pass

    icerik = _icerik_uret(hafta, ctx)

    from actions.code_tools import ensure_workspace_folder
    folder = ensure_workspace_folder("Günlük Planlar")
    stamp = time.strftime("%Y-%m-%d %H.%M")
    out_name = f"{ders_norm.capitalize()} {sinif}.Sınıf {hafta['hafta_no']}.Hafta Günlük Plan {stamp}.docx"
    out_path = folder / out_name
    i = 2
    while out_path.exists():
        out_path = folder / f"{ders_norm.capitalize()} {sinif}.Sınıf {hafta['hafta_no']}.Hafta Günlük Plan {stamp} ({i}).docx"
        i += 1

    try:
        _build_plan_docx(ctx, hafta, icerik, out_path)
    except Exception as e:
        return f"Günlük plan dosyası oluşturulamadı: {e}"

    try:
        _open_file(out_path)
    except Exception:
        pass

    model_adi = (str(get_app_config_value("ollama_model", "")) if get_model_provider() == "ollama" else "Gemini")
    return (f"{ders_info['ad']} — {sinif}. sınıf {hafta['hafta_no']}. hafta günlük planı "
            f"({model_adi} ile üretildi) hazırlandı ve açıldı: {out_path.name}.")
