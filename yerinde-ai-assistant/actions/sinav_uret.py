"""
actions/sinav_uret.py — Kazanımlara dayalı sınav üretimi.

İki farklı kazanım kaynağı modeli:

  BİLİŞİM: İl'in yayınladığı "Konu Soru Dağılım Tablosu" (KSDT / kazanım
  senaryosu) PDF'i var. Bu tablo, her kazanımın hangi sınav + hangi senaryoda
  (1./2. Sınav x 1./2. Senaryo) kaç soru olarak sorulacağını SAYIYLA belirtir.
  Sınav istendiğinde, o sınıf/sınav/senaryo hücresinde sayı bulunan
  kazanımlar birebir kullanılır — kazanım seçimi TAHMİN değil, tablodan.

  ROBOTİK: Böyle bir senaryo tablosu yok (MEB/İl seviyesinde yayınlanmıyor).
  Bunun yerine yıllık plan belgesinin TAMAMI bağlam olarak modele verilir;
  öğretmenin belirttiği konu/ünite kapsamına göre modelin kendisi uygun
  kazanımları seçip soru üretir.

Sonuç, örnek sınavın gözlemlenen biçimine (ortalanmış kalın başlık bloğu,
Adı/Soyadı/Sınıf/Numarası tablosu, "Soru-N:" numaralı sorular, noktalı cevap
alanı) uygun YENİ bir .docx olarak python-docx ile SIFIRDAN inşa edilir —
zümre tutanağının aksine burada soru sayısı her seferinde değiştiği için
mevcut bir dosyayı klonlamak yerine şablonu kod olarak üretiyoruz.
"""

from __future__ import annotations

import json
import platform
import re
import subprocess
import time
from pathlib import Path

import requests

from app_config import get_app_config_value, get_model_provider, ollama_think_value

_IS_WINDOWS = platform.system() == "Windows"
BATCH_SIZE = 5
LLM_TIMEOUT = 600

DERS_BILGI = {
    "bilişim": {"ad": "BİLİŞİM TEKNOLOJİLERİ ve YAZILIM DERSİ", "brans": "Bilişim Teknolojileri ve Yazılım Öğretmeni",
                "kod_onek": "BT"},
    "robotik": {"ad": "SEÇMELİ ROBOTİK KODLAMA DERSİ", "brans": "Bilişim Teknolojileri ve Yazılım Öğretmeni",
                "kod_onek": "RK"},
}


def _tr_lower(s: str) -> str:
    return s.replace("İ", "i").replace("I", "ı").lower()


def _normalize_ders(ders: str) -> str:
    d = _tr_lower(ders or "")
    return "robotik" if "robot" in d else "bilişim"


def _open_file(path: Path) -> None:
    if _IS_WINDOWS:
        import os
        os.startfile(str(path))  # noqa: S606
    else:
        subprocess.Popen(["xdg-open", str(path)],
                         stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)


# ══ 1) Bilişim: KSDT / kazanım senaryosu tablosunu oku ═══════════════════
_KAZANIM_RE = re.compile(r"^([A-ZÇĞİÖŞÜ]{2,4}\.\d+(?:\.\d+){2,3})\.?\s*(.+)$")


def _load_ksdt(path: Path) -> dict:
    """Döner: {sinif_str: [ {kod, metin, s1_sen1, s1_sen2, s2_sen1, s2_sen2}, ... ]}"""
    import pdfplumber
    sonuc: dict[str, list[dict]] = {}
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                sinif = None
                satirlar = []
                for row in table:
                    cells = [(c or "").strip().replace("\n", " ") for c in row]
                    if len(cells) < 2:
                        continue
                    # sınıf tespiti: kazanım kodundan (BT.5. / BT.6.)
                    m = _KAZANIM_RE.match(cells[1])
                    if not m:
                        continue
                    kod = m.group(1)
                    if sinif is None:
                        parts = kod.split(".")
                        sinif = parts[1] if len(parts) > 1 else "?"
                    metin = m.group(2)

                    def _sayi(i):
                        v = cells[i] if i < len(cells) else ""
                        v = re.sub(r"\D", "", v)
                        return int(v) if v else 0

                    satirlar.append({
                        "kod": kod, "metin": metin,
                        "s1_sen1": _sayi(2), "s1_sen2": _sayi(3),
                        "s2_sen1": _sayi(4), "s2_sen2": _sayi(5),
                    })
                if sinif and satirlar:
                    sonuc.setdefault(sinif, []).extend(satirlar)
    return sonuc


def _select_kazanimlar_bilisim(ksdt: dict, sinif: str, sinav_no: str, senaryo_no: str) -> list[dict]:
    satirlar = ksdt.get(str(sinif).strip(), [])
    kolon = f"s{sinav_no}_sen{senaryo_no}"
    secilen = []
    for row in satirlar:
        adet = row.get(kolon, 0)
        if adet and adet > 0:
            secilen.append({"kod": row["kod"], "metin": row["metin"], "adet": adet})
    return secilen


# ══ 2) Robotik: yıllık plandan serbest metin bağlamı ═════════════════════
def _extract_any_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            return "\n".join((p.extract_text() or "") for p in pdf.pages)
    if ext == ".docx":
        import docx
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for t in d.tables:
            for row in t.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(parts)
    if ext in (".xlsx", ".xlsm"):
        import openpyxl
        wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    return path.read_text(encoding="utf-8", errors="replace")


# ══ 3) LLM ile soru üretimi ═══════════════════════════════════════════════
def _extract_json_array(raw: str, expected_len: int | None = None) -> list | None:
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw.strip())
    m = re.search(r"\[.*\]", raw, re.S)
    if not m:
        return None
    try:
        arr = json.loads(m.group(0))
    except Exception:
        return None
    if not isinstance(arr, list):
        return None
    if expected_len is not None and len(arr) != expected_len:
        return None
    return arr


def _call_gemini_text(prompt: str) -> str | None:
    api_key = str(get_app_config_value("gemini_api_key", "") or "").strip()
    if not api_key:
        return None
    try:
        from google import genai
    except ImportError:
        return None
    try:
        client = genai.Client(api_key=api_key)
        for model in ("models/gemini-2.5-flash", "models/gemini-2.0-flash", "models/gemini-2.5-flash-lite"):
            try:
                resp = client.models.generate_content(model=model, contents=[prompt])
                if resp and resp.text:
                    return resp.text
            except Exception:
                continue
    except Exception:
        pass
    return None


def _call_ollama_text(prompt: str, num_predict: int = 2200) -> str | None:
    try:
        host = str(get_app_config_value("ollama_host", "http://localhost:11434") or "http://localhost:11434")
        model = str(get_app_config_value("ollama_model", "llama3.1") or "llama3.1")
        payload = {"model": model, "stream": False,
                   "messages": [{"role": "user", "content": prompt}],
                   "options": {"temperature": 0.5, "num_predict": num_predict},
                   "think": ollama_think_value(), "keep_alive": "30m"}
        r = requests.post(f"{host}/api/chat", json=payload, timeout=LLM_TIMEOUT)
        if r.status_code >= 400:
            payload.pop("think", None)
            r = requests.post(f"{host}/api/chat", json=payload, timeout=LLM_TIMEOUT)
        return r.json().get("message", {}).get("content", "")
    except Exception:
        return None


def _call_llm(prompt: str, num_predict: int = 2200) -> str | None:
    return (_call_gemini_text(prompt) if get_model_provider() == "gemini"
            else _call_ollama_text(prompt, num_predict))


def _questions_prompt_bilisim(kazanimlar: list[dict], ctx: dict) -> str:
    numbered = "\n".join(
        f"{i+1}. [{k['kod']}] {k['metin']} (bu kazanımdan {k['adet']} soru yaz)"
        for i, k in enumerate(kazanimlar)
    )
    toplam_soru = sum(k["adet"] for k in kazanimlar)
    ek = f"\nEK TALİMAT: {ctx['ek_talimat']}" if ctx.get("ek_talimat") else ""
    return (
        f"Sen bir Türk devlet ortaokulunda {ctx['sinif']}. sınıflara "
        f"{ctx['ders_ad']} dersi veren bir öğretmensin. Aşağıdaki kazanımlardan, "
        f"her kazanım için istenen sayıda, {ctx['sinif']}. sınıf seviyesine uygun, "
        f"resmî yazılı sınav üslubunda soru yaz.\n\n"
        f"{_soru_tipi_talimati(ctx)}\n"
        f"Her sorunun cevabı öğrencinin boşluğa elle yazacağı ölçüde kısa "
        f"olmalı (çoktan seçmeli DEĞİL — MEB yazılılarında şık kullanılmaz)."
        f"{ek}\n\n"
        f"KAZANIMLAR:\n{numbered}\n\n"
        f"SADECE şu JSON dizisini döndür, başka HİÇBİR şey yazma: "
        f'[{{"kazanim_kod": "...", "soru_tipi": "klasik|bosluk_doldurma|esleştirme", "soru": "..."}}, ...] '
        f"— toplam tam olarak {toplam_soru} elemanlı olmalı (her kazanım için "
        f"istenen sayı kadar ayrı obje), sıra kazanım sırasına uygun olsun. "
        f"'esleştirme' seçtiysen 'soru' alanına HEM sol liste HEM sağ liste "
        f"maddelerini net biçimde yaz (ör. '1) ... 2) ... A) ... B) ...')."
    )


def _soru_tipi_talimati(ctx: dict) -> str:
    tip = _tr_lower(ctx.get("soru_tipi", "") or "karışık")
    if "klasik" in tip or "açık" in tip:
        return "Tüm sorular KLASİK (açık uçlu, 'nedir/açıklayınız/örnek veriniz') tipte olsun."
    if "boşluk" in tip or "bosluk" in tip:
        return "Tüm sorular BOŞLUK DOLDURMA tipte olsun (cümlede '.....' ile boşluk bırak)."
    if "eşleştirme" in tip or "esleştirme" in tip or "eslestirme" in tip:
        return "Tüm sorular EŞLEŞTİRME tipte olsun (iki liste, öğrenci ilişkilendirsin)."
    return (
        "Soru TİPLERİNİ SEN KARAR VER ve ÇEŞİTLENDİR: çoğunluk (~%60) KLASİK "
        "açık uçlu, bir kısmı (~%25) BOŞLUK DOLDURMA, kalanı (~%15) — özellikle "
        "kavram-tanım eşleştirmeye uygun kazanımlarda — EŞLEŞTİRME olsun. Hangi "
        "kazanımın hangi tipe daha uygun olduğuna kendin karar ver."
    )


def _questions_prompt_robotik(baglam: str, ctx: dict) -> str:
    ek = f"\nEK TALİMAT: {ctx['ek_talimat']}" if ctx.get("ek_talimat") else ""
    kapsam = f"\nKONU/ÜNİTE KAPSAMI: {ctx['konu_kapsam']}" if ctx.get("konu_kapsam") else ""
    n = ctx.get("soru_sayisi") or 8
    # Bağlam çok uzunsa (tüm yıllık plan) kısalt — model bağlam penceresini taşırmasın
    if len(baglam) > 8000:
        baglam = baglam[:8000] + "\n... (yıllık planın devamı kısaltıldı)"
    return (
        f"Sen bir Türk devlet ortaokulunda {ctx['sinif']}. sınıflara "
        f"{ctx['ders_ad']} dersi veren bir öğretmensin. Aşağıda bu dersin YILLIK "
        f"PLANI'ndan alınmış kazanım/konu listesi var. Bu listeden, "
        f"{ctx['sinif']}. sınıf seviyesine uygun, resmî yazılı sınav üslubunda "
        f"TAM OLARAK {n} soru yaz.\n\n"
        f"{_soru_tipi_talimati(ctx)}\n"
        f"Her sorunun cevabı öğrencinin boşluğa elle yazacağı ölçüde kısa "
        f"olmalı (çoktan seçmeli DEĞİL).{kapsam}{ek}\n\n"
        f"YILLIK PLANDAN KAZANIM/KONU BAĞLAMI:\n{baglam}\n\n"
        f"SADECE şu JSON dizisini döndür, başka HİÇBİR şey yazma: "
        f'[{{"kazanim_kod": "", "soru_tipi": "klasik|bosluk_doldurma|esleştirme", "soru": "..."}}, ...] '
        f"— tam olarak {n} elemanlı olmalı (kazanim_kod bulamıyorsan boş bırak, "
        f"önemli olan soru metni)."
    )


def _generate_questions(ctx: dict, kazanimlar: list[dict] | None, robotik_baglam: str | None) -> tuple[list[dict], int]:
    """Döner: (sorular, hata_sayisi). sorular = [{"kazanim_kod","soru"}]"""
    sorular: list[dict] = []
    hata = 0
    if kazanimlar is not None:
        for i in range(0, len(kazanimlar), BATCH_SIZE):
            batch = kazanimlar[i:i + BATCH_SIZE]
            prompt = _questions_prompt_bilisim(batch, ctx)
            beklenen = sum(k["adet"] for k in batch)
            raw = _call_llm(prompt, num_predict=1600)
            arr = _extract_json_array(raw, beklenen) if raw else None
            if not arr:
                raw2 = _call_llm(prompt + "\n\nUNUTMA: SADECE JSON dizisi.", num_predict=1600) if raw else None
                arr = _extract_json_array(raw2, beklenen) if raw2 else None
            if arr:
                sorular.extend(arr)
            else:
                hata += beklenen
                for k in batch:
                    for _ in range(k["adet"]):
                        sorular.append({"kazanim_kod": k["kod"],
                                         "soru": f"{k['metin']} konusunu açıklayınız."})
    else:
        n = ctx.get("soru_sayisi") or 8
        prompt = _questions_prompt_robotik(robotik_baglam or "", ctx)
        raw = _call_llm(prompt, num_predict=1800)
        arr = _extract_json_array(raw, n) if raw else None
        if not arr:
            raw2 = _call_llm(prompt + "\n\nUNUTMA: SADECE JSON dizisi.", num_predict=1800) if raw else None
            arr = _extract_json_array(raw2, n) if raw2 else None
        if arr:
            sorular.extend(arr)
        else:
            hata += n
            sorular.extend([{"kazanim_kod": "", "soru": f"({i+1}. soru üretilemedi — elle ekleyin.)"}
                             for i in range(n)])
    return sorular, hata


# ══ 4) Word belgesini inşa etme ═══════════════════════════════════════════
def _puan_dagit(n: int) -> list[int]:
    if n <= 0:
        return []
    taban = 100 // n
    kalan = 100 - taban * n
    puanlar = [taban] * n
    for i in range(kalan):
        puanlar[n - 1 - i] += 1
    return puanlar


def _build_exam_docx(ctx: dict, sorular: list[dict], out_path: Path) -> None:
    import docx
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    d = docx.Document()
    section = d.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Cm(1.5))

    def _p(text="", bold=False, italic=False, size=12, align=None, space_after=6):
        p = d.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            r = p.add_run(text)
            r.bold = bold
            r.italic = italic
            r.font.size = Pt(size)
        return p

    def _bottom_border(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    okul = ctx.get("okul_adi") or "[OKUL ADI]"
    _p(f"{ctx['ders_yili']} EĞİTİM ÖĞRETİM YILI {okul}",
       bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _p(f"{ctx['ders_ad']}", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _p(f"{ctx['sinif']}.SINIFLAR {ctx['donem']}.DÖNEM {ctx['sinav_no']}.YAZILI SINAVI",
       bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    puanlar = _puan_dagit(len(sorular))
    tekduzey = len(set(puanlar)) <= 1
    if tekduzey and puanlar:
        _p(f"Her soru {puanlar[0]} puandır.", bold=True, size=11, space_after=10)
    else:
        _p("Soru puanları sorunun yanında belirtilmiştir.", bold=True, size=11, space_after=10)

    # Öğrenci bilgi tablosu
    tbl = d.add_table(rows=2, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    hdr = ["ADI", "SOYADI", "SINIF", "NUMARASI"]
    for i, h in enumerate(hdr):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(4):
        tbl.rows[1].cells[i].text = ""
    d.add_paragraph().paragraph_format.space_after = Pt(12)

    for i, soru in enumerate(sorular, start=1):
        p = _p()
        r1 = p.add_run(f"Soru-{i}: ")
        r1.bold = True
        r1.font.size = Pt(11)
        puan_txt = f" ({puanlar[i-1]} Puan)" if not tekduzey else ""
        kod_txt = f" ({soru.get('kazanim_kod')})" if soru.get("kazanim_kod") else ""
        r2 = p.add_run(f"{soru.get('soru', '').strip()}{puan_txt}{kod_txt}")
        r2.font.size = Pt(11)

        for _ in range(2):
            ans = _p("…" * 110, size=11, space_after=2)
        d.paragraphs[-1].paragraph_format.space_after = Pt(10)
        _bottom_border(d.paragraphs[-1])

    _p("")
    _p("Sınav süresi 1 ders saatidir.     Başarılar dilerim. 🙂",
       align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=4)
    if ctx.get("ogretmen_adi"):
        _p(ctx["ogretmen_adi"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=0)
    _p(ctx.get("brans", ""), align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=0)

    d.save(str(out_path))


# ══ 5) Genel giriş noktası ════════════════════════════════════════════════
def sinav_olustur(
    ders: str = "bilişim",
    sinif: str = "5",
    sinav_no: str = "1",
    senaryo_no: str = "1",
    donem: str = "2",
    soru_sayisi: int = 0,
    konu_kapsam: str = "",
    ek_talimat: str = "",
    ksdt_dosya_yolu: str = "",
    yillik_plan_dosya_yolu: str = "",
    soru_tipi: str = "karışık",
) -> str:
    from actions.belge_referanslari import referans_yolu
    from actions.zumre_tutanagi import _read_reference as _read_zumre_ref  # okul adı / öğretmen adı için

    ders_norm = _normalize_ders(ders)
    sinif = re.sub(r"\D", "", str(sinif)) or "5"
    sinav_no = re.sub(r"\D", "", str(sinav_no)) or "1"
    senaryo_no = re.sub(r"\D", "", str(senaryo_no)) or "1"
    donem = re.sub(r"\D", "", str(donem)) or "2"

    ders_info = DERS_BILGI[ders_norm]
    ctx = {
        "sinif": sinif, "donem": donem, "sinav_no": sinav_no,
        "ders_ad": ders_info["ad"], "brans": ders_info["brans"],
        "ders_yili": "2026-2027", "ek_talimat": ek_talimat.strip(),
        "konu_kapsam": konu_kapsam.strip(), "soru_sayisi": int(soru_sayisi or 0),
        "okul_adi": "", "ogretmen_adi": "", "soru_tipi": soru_tipi.strip(),
    }

    # Okul adı / öğretmen adını, kayıtlı zümre örneğinden ödünç al (varsa) —
    # aynı bilgiyi tekrar tekrar sormaya gerek kalmasın.
    try:
        zumre_key = "robotik_zumre_ornek" if ders_norm == "robotik" else "bilisim_zumre_ornek"
        zref = referans_yolu(zumre_key)
        if zref:
            parsed = _read_zumre_ref(zref)
            ctx["ogretmen_adi"] = ""
            title_words = parsed.get("zumre_adi", "")
            # okul adı, referans dosyasındaki ilk ortalanmış başlık paragrafından
            import docx as _docx
            dd = _docx.Document(str(zref))
            if dd.paragraphs and dd.paragraphs[0].text.strip():
                ctx["okul_adi"] = dd.paragraphs[0].text.strip()
            baskan = parsed.get("info", {}).get(
                "Zümre Başkanı Adı Soyadı, Telefonu, \nE-Posta Adresi", "") or \
                parsed.get("info", {}).get("Zümre Başkanı Adı Soyadı, Telefonu,", "")
            if baskan:
                ctx["ogretmen_adi"] = baskan.split("0 ")[0].strip() if "0 " in baskan else baskan.split("(")[0].strip()
    except Exception:
        pass

    if ders_norm == "bilişim":
        ksdt_path = (Path(ksdt_dosya_yolu) if ksdt_dosya_yolu and Path(ksdt_dosya_yolu).exists()
                     else referans_yolu("bilisim_kazanim_senaryosu"))
        if not ksdt_path:
            return ("Bilişim kazanım senaryosu (KSDT) tablosu bulamadım. Önce İl'in yayınladığı "
                     "'Konu Soru Dağılım Tablosu' PDF'ini 'DOSYA YÜKLE' ile yükleyip 'bunu bilişim "
                     "kazanım senaryosu olarak kaydet' de, sonra tekrar iste.")
        try:
            ksdt = _load_ksdt(ksdt_path)
        except Exception as e:
            return f"Kazanım senaryosu tablosu okunamadı: {e}"
        kazanimlar = _select_kazanimlar_bilisim(ksdt, sinif, sinav_no, senaryo_no)
        if not kazanimlar:
            return (f"{sinif}. sınıf, {sinav_no}. sınav, {senaryo_no}. senaryo için kazanım "
                     f"tabloda bulunamadı — sınıf/sınav/senaryo numarasını kontrol eder misin?")
        sorular, hata = _generate_questions(ctx, kazanimlar, None)
    else:  # robotik
        yp_path = (Path(yillik_plan_dosya_yolu) if yillik_plan_dosya_yolu and Path(yillik_plan_dosya_yolu).exists()
                   else referans_yolu("robotik_yillik_plan"))
        if not yp_path:
            return ("Robotik yıllık planını bulamadım. Önce 'DOSYA YÜKLE' ile yükleyip 'bunu robotik "
                     "yıllık planı olarak kaydet' de, sonra tekrar iste.")
        try:
            baglam = _extract_any_text(yp_path)
        except Exception as e:
            return f"Yıllık plan okunamadı: {e}"
        if not ctx["soru_sayisi"]:
            ctx["soru_sayisi"] = 8
        sorular, hata = _generate_questions(ctx, None, baglam)

    if not sorular:
        return "Soru üretilemedi — model yanıt vermedi. Ollama/Gemini bağlantısını kontrol eder misin?"

    from actions.code_tools import ensure_workspace_folder
    folder = ensure_workspace_folder("Sınavlar")
    stamp = time.strftime("%Y-%m-%d %H.%M")
    out_name = f"{ders_norm.capitalize()} {sinif}.Sınıf {donem}.Dönem {sinav_no}.Sınav {stamp}.docx"
    out_path = folder / out_name
    i = 2
    while out_path.exists():
        out_path = folder / f"{ders_norm.capitalize()} {sinif}.Sınıf {donem}.Dönem {sinav_no}.Sınav {stamp} ({i}).docx"
        i += 1

    try:
        _build_exam_docx(ctx, sorular, out_path)
    except Exception as e:
        return f"Sınav dosyası oluşturulamadı: {e}"

    try:
        _open_file(out_path)
    except Exception:
        pass

    model_adi = (str(get_app_config_value("ollama_model", "")) if get_model_provider() == "ollama" else "Gemini")
    uyari = f" NOT: {hata} soru üretilemediği için yer tutucu bırakıldı, elle gözden geçir." if hata else ""
    return (f"{ders_info['ad']} — {sinif}. sınıf {sinav_no}. sınav ({len(sorular)} soru, {model_adi} ile "
            f"üretildi) hazırlandı ve açıldı: {out_path.name}.{uyari}")
